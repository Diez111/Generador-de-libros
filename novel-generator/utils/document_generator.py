from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.enum.section import WD_ORIENT
import regex

class DocumentGenerator:
    @staticmethod
    def generar_word(context: dict):
        doc = Document()
        DocumentGenerator.aplicar_estilos(doc)
        DocumentGenerator.agregar_portada(doc, context)
        DocumentGenerator.agregar_indice(doc, context)
        
        for cap in context['estado']['trama']['capitulos']:
            DocumentGenerator.agregar_capitulo(doc, cap)
        
        nombre = f"{DocumentGenerator.nombre_seguro(context['meta']['titulo'])}.docx"
        doc.save(nombre)
        return nombre

    @staticmethod
    def aplicar_estilos(doc):
        styles = doc.styles
        
        titulo_style = styles.add_style('TituloLibro', WD_STYLE_TYPE.PARAGRAPH)
        titulo_style.font.name = 'Times New Roman'
        titulo_style.font.size = Pt(28)
        titulo_style.font.bold = True
        titulo_style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        titulo_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        titulo_style.paragraph_format.space_after = Pt(24)
        
        autor_style = styles.add_style('AutorLibro', WD_STYLE_TYPE.PARAGRAPH)
        autor_style.font.name = 'Times New Roman'
        autor_style.font.size = Pt(14)
        autor_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        autor_style.paragraph_format.space_after = Pt(60)
        
        capitulo_style = styles.add_style('TituloCapitulo', WD_STYLE_TYPE.PARAGRAPH)
        capitulo_style.font.name = 'Times New Roman'
        capitulo_style.font.size = Pt(16)
        capitulo_style.font.bold = True
        capitulo_style.paragraph_format.page_break_before = True
        capitulo_style.paragraph_format.space_after = Pt(12)
        
        cuerpo_style = styles.add_style('CuerpoTexto', WD_STYLE_TYPE.PARAGRAPH)
        cuerpo_style.font.name = 'Times New Roman'
        cuerpo_style.font.size = Pt(12)
        cuerpo_style.paragraph_format.line_spacing = 1.5
        cuerpo_style.paragraph_format.space_after = Pt(6)
        cuerpo_style.paragraph_format.first_line_indent = Inches(0.3)
        
        section = doc.sections[0]
        section.page_width = Inches(6)
        section.page_height = Inches(9)
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    @staticmethod
    def agregar_portada(doc, context):
        p = doc.add_paragraph(style='TituloLibro')
        p.add_run(context['meta']['titulo'].upper())
        
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        
        p = doc.add_paragraph(style='AutorLibro')
        p.add_run(f"POR {context['meta']['autor'].upper()}")

    @staticmethod
    def agregar_indice(doc, context):
        doc.add_paragraph("ÍNDICE", style='TituloCapitulo')
        for cap in context['estado']['trama']['capitulos']:
            doc.add_paragraph(
                f"Capítulo {cap['num']}: {cap['titulo']}", 
                style='CuerpoTexto'
            )
        doc.add_page_break()

    @staticmethod
    def agregar_capitulo(doc, cap):
        p = doc.add_paragraph(style='TituloCapitulo')
        p.add_run(f"Capítulo {cap['num']}\n{cap['titulo']}".upper())
        
        for parrafo in cap['contenido'].split('\n\n'):
            if parrafo.strip():
                p = doc.add_paragraph(style='CuerpoTexto')
                p.add_run(parrafo.strip())
        
        doc.add_page_break()

    @staticmethod
    def nombre_seguro(texto: str) -> str:
        return regex.sub(r'[\\/*?:"<>|]', "", texto)[:45].strip()